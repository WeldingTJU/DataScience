import openai
import logging
import time
import os
from dotenv import load_dotenv
from docx import Document

load_dotenv()

DEEPSEEK_API_KEY = os.getenv('DEEPSEEK_API_KEY')
if not DEEPSEEK_API_KEY:
    logging.error("DeepSeek API key is not set. Please define DEEPSEEK_API_KEY in your environment variables.")
    exit("DeepSeek API key is not set. Please define DEEPSEEK_API_KEY in your environment variables.")
else:
    openai.api_key = DEEPSEEK_API_KEY
    openai.api_base = "https://api.deepseek.com/v1"  # DeepSeek API endpoint

logging.basicConfig(
    filename='process.log',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

SYSTEM_MESSAGE_CONTENT = """
You are tasked with extracting information about fatigue test specimens.
For each document provided, extract detailed fatigue specimen information, specifically including:
**Parameter**:
1. **Title**: Title of the document where the specimen is described
2. **Base Material**: Material of the base metal in the fatigue specimen
3. **Base Young's Modulus (GPa)**: Young's modulus of the base material, unit in GPa
4. **Base Yield Strength (MPa)**: Yield strength of the base material, unit in MPa
5. **Base Ultimate Tensile Strength (MPa)**: Ultimate tensile strength of the base material, unit in MPa
6. **Base Elongation (%)**: Elongation of the base material, in percentage
7. **Weld Material**: Material of the weld in the fatigue specimen
8. **Weld Young's Modulus (GPa)**: Young's modulus of the weld material, unit in GPa
9. **Weld Yield Strength (MPa)**: Yield strength of the weld material, unit in MPa
10. **Weld Ultimate Tensile Strength (MPa)**: Ultimate tensile strength of the weld material, unit in MPa
11. **Weld Elongation (%)**: Elongation of the weld material, in percentage
12. **Processing**: Indicates if the fatigue specimen underwent processing, e.g., Annealing, Quenching, Tempering, Ultrasonic Impact Treatment
13. **Type of welding joint**: e.g., Butt Joint, T-Joint
14. **Welding method**: e.g., TIG, FSW, MIG
15. **Welding voltage (V)**: Unit in V
16. **Welding current (A)**: Unit in A
17. **Welding speed (mm/s)**: Unit in mm/s
18. **Preheat temperature of Welding (°C)**: Unit in °C
19. **Welding environment**
20. **Residual Stress**: Corresponding to the fatigue specimen
21. **Type of fatigue specimen**: e.g., full-scale, standard, notched, round bar, wide plate, compact tension (CT), single edge notch bend (SENB), dog bone, tubular, cantilever beam, sheet,flat
22. **Thickness (mm) of fatigue specimen**: Thickness of fatigue specimen, unit in mm
23. **Types of fatigue tests**: e.g., Axial, Torsional, Multiaxial, Rotating Bending
24. **Fatigue test temperature (°C)**: Unit in °C
25. **Fatigue test environment**
26. **Load ratio of fatigue test**
27. **Frequency (Hz) of fatigue test**: Unit in Hz
28. **Fatigue test machine**
29. **Fatigue test standard**
30. **Fatigue test control**: e.g., Stress, Load, Displacement
31. **Stress concentration factor**
32. **Failure location**: Failure location in fatigue test, e.g., weld toe,weld root.

**S-N**:
- **Life N (cycles)**: The number of loading cycles at a specific stress amplitude
- **Stress amplitude σa (MPa)**: The stress amplitude applied in the fatigue test, Unit in MPa
- **Runout**: Use "1" if this specimen mentions runout, otherwise indicate with "-"

**e-N**:
- **Life N (cycles)**: The number of loading cycles at a specific strain amplitude
- **Strain amplitude εa (mm/mm)**: The strain amplitude applied in the fatigue test. Unit in mm/mm
- **Runout**: Use "1" if this specimen mentions runout, otherwise indicate with "-"

Important:
1. For S-N and e-N data, if they are difficult to identify in the table, please indicate where the S-N and e-N data are located
2. Do not fabricate any specimen information; if certain information is not found in the document, use "-"
3. Ensure that all fields have the appropriate units where required
4. Ensure the data is accurate
"""

class DeepSeekHandler:

    @staticmethod
    def extract_information(file_content, max_retries=3, delay=1):
        for attempt in range(max_retries):
            try:
                response = openai.ChatCompletion.create(
                    model="deepseek-chat",
                    messages=[
                        {"role": "system", "content": SYSTEM_MESSAGE_CONTENT},
                        {"role": "user", "content": file_content}
                    ],
                    max_tokens=6000,
                    temperature=0
                )
                fatigue_data = response.choices[0].message['content'].strip()
                logging.info("Successfully called DeepSeek to extract information")
                return fatigue_data
            except openai.error.RateLimitError as e:
                logging.error(f"API rate limit exceeded (attempt {attempt + 1}), error: {e}")
                time.sleep(delay * (2 ** attempt))
            except openai.error.APIError as e:
                logging.error(f"API error (attempt {attempt + 1}): {e}")
                time.sleep(delay)
            except Exception as e:
                logging.error(f"Unexpected error: {e}")
                return None
        return None

def read_word_file(file_path):
    doc = Document(file_path)
    full_text = []

    for para in doc.paragraphs:
        full_text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text)
            full_text.append("\t".join(row_text))

    return "\n".join(full_text)

def read_text_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def process_all_files_in_directory(directory_path, output_directory):
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)

    for filename in os.listdir(directory_path):
        file_path = os.path.join(directory_path, filename)

        if os.path.isfile(file_path):
            # Select reading method based on file extension
            if filename.endswith('.docx'):
                file_content = read_word_file(file_path)
            elif filename.endswith('.txt'):
                file_content = read_text_file(file_path)
            else:
                logging.warning(f"Unsupported file type: {filename}")
                continue

            extracted_data = DeepSeekHandler.extract_information(file_content)

            if extracted_data:
                output_filename = os.path.splitext(filename)[0] + ".txt"
                output_file_path = os.path.join(output_directory, output_filename)
                with open(output_file_path, "w", encoding="utf-8") as output_file:
                    output_file.write(extracted_data)
                logging.info(f"Successfully saved to file: {output_file_path}")
        else:
            logging.warning(f"Skipping directory: {file_path}")

if __name__ == "__main__":
    directory_path = r"F:\unprocessed"
    output_directory = r"F:\processed"
    process_all_files_in_directory(directory_path, output_directory)
